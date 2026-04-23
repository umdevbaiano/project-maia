"""
Maia Platform — DOCX Generator (ABNT NBR 14724)
Converts markdown text to formatted .docx documents with full legal/ABNT formatting.
"""
import io
import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _add_page_number_footer(section, workspace_name: str = ""):
    """Add a footer with workspace name on the left and 'Page X of Y' on the right."""
    footer = section.footer
    footer.is_linked_to_previous = False

    # Clear default footer paragraph and set up a two-column layout with a tab
    p = footer.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(6)

    # Office name on left
    if workspace_name:
        run_office = p.add_run(workspace_name)
        run_office.font.name = "Times New Roman"
        run_office.font.size = Pt(10)
        run_office.font.color.rgb = RGBColor(100, 100, 100)

    # Tab to right-align page number
    tab_stop = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9072")  # right margin position
    tab_stop.append(tab)
    p._p.get_or_add_pPr().append(tab_stop)
    p.add_run("\t")

    # "Página " text
    run_pg = p.add_run("Página ")
    run_pg.font.name = "Times New Roman"
    run_pg.font.size = Pt(10)
    run_pg.font.color.rgb = RGBColor(100, 100, 100)

    # Current page number field
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    run_num = p.add_run()
    run_num._r.append(fld_char)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run_instr = p.add_run()
    run_instr._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run_end = p.add_run()
    run_end._r.append(fld_char_end)
    for r in [run_num, run_instr, run_end]:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)

    # " de " and total pages
    run_de = p.add_run(" de ")
    run_de.font.name = "Times New Roman"
    run_de.font.size = Pt(10)
    run_de.font.color.rgb = RGBColor(100, 100, 100)

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "begin")
    run_total = p.add_run()
    run_total._r.append(fld_char2)

    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = " NUMPAGES "
    run_instr2 = p.add_run()
    run_instr2._r.append(instr2)

    fld_char2_end = OxmlElement("w:fldChar")
    fld_char2_end.set(qn("w:fldCharType"), "end")
    run_end2 = p.add_run()
    run_end2._r.append(fld_char2_end)
    for r in [run_total, run_instr2, run_end2]:
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(100, 100, 100)


def _add_header(section, titulo: str, workspace_name: str = ""):
    """Add a header with the document title and a separator line."""
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if workspace_name:
        run_ws = p.add_run(workspace_name.upper())
        run_ws.font.name = "Times New Roman"
        run_ws.font.size = Pt(9)
        run_ws.bold = True
        run_ws.font.color.rgb = RGBColor(80, 80, 80)
        p.add_run(" · ")

    # Truncate title for header
    short_title = titulo[:80] + ("…" if len(titulo) > 80 else "")
    run_title = p.add_run(short_title)
    run_title.font.name = "Times New Roman"
    run_title.font.size = Pt(9)
    run_title.font.color.rgb = RGBColor(100, 100, 100)
    run_title.italic = True

    # Bottom border on the header paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


def markdown_to_docx(
    markdown_text: str,
    titulo: str = "Documento Jurídico",
    workspace_name: str = "",
) -> io.BytesIO:
    """
    Convert markdown text to a fully formatted .docx document with ABNT NBR 14724 formatting.
    Returns a BytesIO buffer ready for download.
    """
    doc = Document()

    # ── Page layout (ABNT NBR 14724: 3cm top/left, 2cm bottom/right) ──
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        _add_header(section, titulo, workspace_name)
        _add_page_number_footer(section, workspace_name)

    # ── Default body style ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    style.paragraph_format.line_spacing = Pt(18)   # 1.5 × 12pt = 18pt
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)

    # ── Parse markdown line by line ──
    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line:
            i += 1
            continue

        # H1 — centred, bold, uppercase, 14pt
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip().upper())
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Times New Roman"
            i += 1
            continue

        # H2 — left, bold, uppercase, 13pt
        if line.startswith("## "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(line[3:].strip().upper())
            run.bold = True
            run.font.size = Pt(13)
            run.font.name = "Times New Roman"
            i += 1
            continue

        # H3 — left, bold, 12pt
        if line.startswith("### "):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(line[4:].strip())
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "CCCCCC")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Blockquote — indented, italic, 11pt (legal citations)
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.left_indent = Cm(4)
            p.paragraph_format.right_indent = Cm(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s", line):
            text = re.sub(r"^\d+\.\s", "", line).strip()
            p = doc.add_paragraph(style="List Number")
            _add_formatted_run(p, text)
            i += 1
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p = doc.add_paragraph(style="List Bullet")
            _add_formatted_run(p, text)
            i += 1
            continue

        # Markdown table
        if "|" in line and i + 1 < len(lines) and "---" in lines[i + 1]:
            rows = []
            while i < len(lines) and "|" in lines[i]:
                cells = [c.strip() for c in lines[i].strip("|").split("|")]
                if not all(c.replace("-", "").replace(":", "") == "" for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                _add_table(doc, rows)
            continue

        # Regular paragraph — justified, first-line indent 1.25cm (ABNT)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(18)
        _add_formatted_run(p, line)
        i += 1

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _add_formatted_run(paragraph, text: str):
    """Parse inline markdown formatting (bold, italic) and add styled runs."""
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Times New Roman"
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = "Times New Roman"
        else:
            run = paragraph.add_run(part)
            run.font.name = "Times New Roman"


def _add_table(doc, rows: list):
    """Add a formatted table to the document."""
    if not rows:
        return
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text.strip()
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(10)
                        if i == 0:
                            run.bold = True
